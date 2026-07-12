from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.ingestion.exceptions import (
    CorruptedDocumentError,
    DocumentNotFoundError,
    DocumentTooLargeError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from app.ingestion.loaders.base import DocumentLoader
from app.ingestion.normalizer import normalize_text
from app.schemas.document import DocumentSection, LoadedDocument


class DOCXDocumentLoader(DocumentLoader):
    """Extract heading-aware sections from Microsoft Word documents."""

    supported_extensions = {".docx"}

    def __init__(self, max_file_size_bytes: int = 20 * 1024 * 1024) -> None:
        if max_file_size_bytes <= 0:
            raise ValueError("max_file_size_bytes must be greater than zero.")

        self.max_file_size_bytes = max_file_size_bytes

    def load(self, file_path: Path) -> LoadedDocument:
        """Load a DOCX document while preserving heading context."""
        resolved_path = file_path.resolve()

        if not resolved_path.exists():
            raise DocumentNotFoundError(f"Document does not exist: {resolved_path}")

        if not resolved_path.is_file():
            raise DocumentNotFoundError(f"Document path is not a file: {resolved_path}")

        extension = resolved_path.suffix.lower()

        if extension not in self.supported_extensions:
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {extension or 'no extension'}"
            )

        file_size = resolved_path.stat().st_size

        if file_size > self.max_file_size_bytes:
            raise DocumentTooLargeError(
                "Document exceeds the configured file-size limit."
            )

        try:
            document = Document(resolved_path)
        except (
            BadZipFile,
            PackageNotFoundError,
            ValueError,
            KeyError,
        ) as error:
            raise CorruptedDocumentError(
                "DOCX document could not be parsed."
            ) from error

        sections: list[DocumentSection] = []
        current_heading: str | None = None
        current_paragraphs: list[str] = []

        def save_current_section() -> None:
            if not current_paragraphs:
                return

            content = normalize_text("\n\n".join(current_paragraphs))

            if not content:
                return

            sections.append(
                DocumentSection(
                    section_index=len(sections),
                    content=content,
                    page_number=None,
                    heading=current_heading,
                )
            )

            current_paragraphs.clear()

        for paragraph in document.paragraphs:
            paragraph_text = normalize_text(paragraph.text)

            if not paragraph_text:
                continue

            style_name = paragraph.style.name if paragraph.style is not None else ""

            if style_name.lower().startswith("heading"):
                save_current_section()
                current_heading = paragraph_text
                continue

            current_paragraphs.append(paragraph_text)

        save_current_section()

        if not sections and current_heading:
            sections.append(
                DocumentSection(
                    section_index=0,
                    content=current_heading,
                    page_number=None,
                    heading=current_heading,
                )
            )

        if not sections:
            raise EmptyDocumentError("DOCX document contains no usable text.")

        combined_parts: list[str] = []

        for section in sections:
            if section.heading:
                combined_parts.append(f"{section.heading}\n\n{section.content}")
            else:
                combined_parts.append(section.content)

        content = "\n\n".join(combined_parts)

        return LoadedDocument(
            file_name=resolved_path.name,
            file_extension=extension,
            source_path=resolved_path,
            content=content,
            character_count=len(content),
            word_count=len(content.split()),
            sections=sections,
        )
