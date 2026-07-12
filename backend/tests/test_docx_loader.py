from pathlib import Path

import pytest
from docx import Document

from app.ingestion.exceptions import (
    CorruptedDocumentError,
    DocumentTooLargeError,
    EmptyDocumentError,
    UnsupportedFileTypeError,
)
from app.ingestion.loaders.docx import DOCXDocumentLoader


def create_heading_document(path: Path) -> None:
    document = Document()

    document.add_heading("Remote Work", level=1)
    document.add_paragraph("Employees may work remotely for three days per week.")

    document.add_heading("Paid Leave", level=1)
    document.add_paragraph("Employees receive eighteen paid leave days.")

    document.save(path)


def test_docx_loader_extracts_heading_aware_sections(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "employee_handbook.docx"
    create_heading_document(document_path)

    loader = DOCXDocumentLoader()
    document = loader.load(document_path)

    assert document.file_name == "employee_handbook.docx"
    assert document.file_extension == ".docx"
    assert document.source_path == document_path.resolve()
    assert len(document.sections) == 2

    first_section = document.sections[0]
    second_section = document.sections[1]

    assert first_section.section_index == 0
    assert first_section.page_number is None
    assert first_section.heading == "Remote Work"
    assert first_section.content == (
        "Employees may work remotely for three days per week."
    )

    assert second_section.section_index == 1
    assert second_section.page_number is None
    assert second_section.heading == "Paid Leave"
    assert second_section.content == ("Employees receive eighteen paid leave days.")


def test_docx_loader_handles_document_without_headings(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "notes.docx"

    source_document = Document()
    source_document.add_paragraph("First policy paragraph.")
    source_document.add_paragraph("Second policy paragraph.")
    source_document.save(document_path)

    loader = DOCXDocumentLoader()
    document = loader.load(document_path)

    assert len(document.sections) == 1
    assert document.sections[0].heading is None
    assert document.sections[0].content == (
        "First policy paragraph.\n\nSecond policy paragraph."
    )


def test_docx_loader_rejects_empty_document(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "empty.docx"
    Document().save(document_path)

    loader = DOCXDocumentLoader()

    with pytest.raises(
        EmptyDocumentError,
        match="no usable text",
    ):
        loader.load(document_path)


def test_docx_loader_rejects_corrupted_document(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "broken.docx"
    document_path.write_bytes(b"not a real docx file")

    loader = DOCXDocumentLoader()

    with pytest.raises(
        CorruptedDocumentError,
        match="could not be parsed",
    ):
        loader.load(document_path)


def test_docx_loader_rejects_unsupported_extension(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "notes.txt"
    document_path.write_text("notes", encoding="utf-8")

    loader = DOCXDocumentLoader()

    with pytest.raises(
        UnsupportedFileTypeError,
        match="Unsupported file type",
    ):
        loader.load(document_path)


def test_docx_loader_rejects_large_document(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "large.docx"
    document_path.write_bytes(b"123456")

    loader = DOCXDocumentLoader(max_file_size_bytes=5)

    with pytest.raises(
        DocumentTooLargeError,
        match="file-size limit",
    ):
        loader.load(document_path)


def test_docx_loader_rejects_invalid_size_configuration() -> None:
    with pytest.raises(
        ValueError,
        match="max_file_size_bytes must be greater than zero",
    ):
        DOCXDocumentLoader(max_file_size_bytes=0)
