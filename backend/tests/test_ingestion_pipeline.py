from pathlib import Path

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from app.ingestion.pipeline import DocumentIngestionPipeline


def test_pipeline_loads_and_chunks_document(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "policy.txt"
    document_path.write_text(
        "zero one two three four five six seven eight nine",
        encoding="utf-8",
    )

    pipeline = DocumentIngestionPipeline(
        chunk_size=5,
        chunk_overlap=2,
    )

    result = pipeline.ingest(document_path)

    assert result.document.file_name == "policy.txt"
    assert result.document.word_count == 10
    assert result.chunk_count == 3

    assert [chunk.text for chunk in result.chunks] == [
        "zero one two three four",
        "three four five six seven",
        "six seven eight nine",
    ]


def test_pipeline_uses_default_text_loader(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "guide.md"
    document_path.write_text(
        "# Guide\n\nRun the health endpoint.",
        encoding="utf-8",
    )

    pipeline = DocumentIngestionPipeline(
        chunk_size=20,
        chunk_overlap=5,
    )

    result = pipeline.ingest(document_path)

    assert result.document.file_extension == ".md"
    assert result.chunk_count == 1
    assert result.chunks[0].text == ("# Guide Run the health endpoint.")


@pytest.mark.parametrize(
    ("chunk_size", "chunk_overlap", "expected_message"),
    [
        (0, 0, "chunk_size must be greater than zero."),
        (-1, 0, "chunk_size must be greater than zero."),
        (5, -1, "chunk_overlap cannot be negative."),
        (5, 5, "chunk_overlap must be smaller than chunk_size."),
        (5, 6, "chunk_overlap must be smaller than chunk_size."),
    ],
)
def test_pipeline_rejects_invalid_chunk_configuration(
    chunk_size: int,
    chunk_overlap: int,
    expected_message: str,
) -> None:
    with pytest.raises(ValueError, match=expected_message):
        DocumentIngestionPipeline(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )


def test_pipeline_preserves_pdf_page_numbers(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "policies.pdf"

    pdf = canvas.Canvas(str(document_path))
    pdf.drawString(72, 720, "Remote work policy applies.")
    pdf.showPage()
    pdf.drawString(72, 720, "Paid leave policy applies.")
    pdf.showPage()
    pdf.save()

    pipeline = DocumentIngestionPipeline(
        chunk_size=20,
        chunk_overlap=5,
    )

    result = pipeline.ingest(document_path)

    assert result.document.file_extension == ".pdf"
    assert result.chunk_count == 2

    assert result.chunks[0].chunk_index == 0
    assert result.chunks[0].section_index == 0
    assert result.chunks[0].page_number == 1
    assert result.chunks[0].text == "Remote work policy applies."

    assert result.chunks[1].chunk_index == 1
    assert result.chunks[1].section_index == 1
    assert result.chunks[1].page_number == 2
    assert result.chunks[1].text == "Paid leave policy applies."


def test_pipeline_preserves_docx_headings(
    tmp_path: Path,
) -> None:
    document_path = tmp_path / "policies.docx"

    source_document = Document()
    source_document.add_heading("Remote Work", level=1)
    source_document.add_paragraph("Employees may work remotely three days per week.")
    source_document.add_heading("Paid Leave", level=1)
    source_document.add_paragraph("Employees receive eighteen paid leave days.")
    source_document.save(document_path)

    pipeline = DocumentIngestionPipeline(
        chunk_size=20,
        chunk_overlap=5,
    )

    result = pipeline.ingest(document_path)

    assert result.chunk_count == 2

    assert result.chunks[0].section_index == 0
    assert result.chunks[0].page_number is None
    assert result.chunks[0].heading == "Remote Work"

    assert result.chunks[1].section_index == 1
    assert result.chunks[1].page_number is None
    assert result.chunks[1].heading == "Paid Leave"
