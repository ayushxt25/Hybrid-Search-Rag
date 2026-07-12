from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from app.main import app

client = TestClient(app)


def test_ingest_document_processes_valid_text_file() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "remote_policy.txt",
                (
                    b"Remote Work Policy\n\n"
                    b"Employees may work remotely for three days per week."
                ),
                "text/plain",
            )
        },
    )

    assert response.status_code == 200

    body = response.json()

    assert body["status"] == "processed"
    assert body["file_name"] == "remote_policy.txt"
    assert body["file_extension"] == ".txt"
    assert body["word_count"] == 12
    assert body["chunk_count"] == 1
    assert len(body["chunks"]) == 1
    assert body["chunks"][0]["chunk_index"] == 0
    assert body["chunks"][0]["section_index"] == 0
    assert body["chunks"][0]["page_number"] is None
    assert body["chunks"][0]["heading"] is None


def test_ingest_document_processes_markdown_file() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "guide.md",
                b"# Guide\n\nRun the health endpoint.",
                "text/markdown",
            )
        },
    )

    assert response.status_code == 200
    assert response.json()["file_extension"] == ".md"


def test_ingest_document_rejects_unsupported_file_type() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "employees.csv",
                b"name,department",
                "text/csv",
            )
        },
    )

    assert response.status_code == 415
    assert response.json() == {
        "detail": ("Only .txt, .md, .pdf and .docx files are currently supported.")
    }


def test_ingest_document_rejects_empty_file() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "empty.txt",
                b"",
                "text/plain",
            )
        },
    )

    assert response.status_code == 422
    assert response.json() == {"detail": "Uploaded document is empty."}


def test_ingest_document_rejects_non_utf8_file() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "invalid.txt",
                b"\xff\xfe\xfa",
                "text/plain",
            )
        },
    )

    assert response.status_code == 422
    assert response.json() == {"detail": "Document must use UTF-8 encoding."}


def test_ingest_document_requires_file() -> None:
    response = client.post("/api/v1/documents/ingest")

    assert response.status_code == 422


def create_pdf_bytes() -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)

    pdf.drawString(
        72,
        720,
        "Remote work is allowed for three days per week.",
    )
    pdf.showPage()

    pdf.drawString(
        72,
        720,
        "Employees receive eighteen paid leave days.",
    )
    pdf.showPage()

    pdf.save()

    return buffer.getvalue()


def test_ingest_document_processes_pdf_with_page_metadata() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "employee_handbook.pdf",
                create_pdf_bytes(),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 200

    body = response.json()

    assert body["status"] == "processed"
    assert body["file_name"] == "employee_handbook.pdf"
    assert body["file_extension"] == ".pdf"
    assert body["chunk_count"] == 2

    assert body["chunks"][0]["section_index"] == 0
    assert body["chunks"][0]["page_number"] == 1
    assert body["chunks"][0]["text"] == (
        "Remote work is allowed for three days per week."
    )

    assert body["chunks"][1]["section_index"] == 1
    assert body["chunks"][1]["page_number"] == 2
    assert body["chunks"][1]["text"] == ("Employees receive eighteen paid leave days.")


def test_ingest_document_rejects_pdf_without_text() -> None:
    buffer = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(buffer)

    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "scanned.pdf",
                buffer.getvalue(),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 422
    assert response.json() == {
        "detail": ("PDF contains no extractable text and may require OCR.")
    }


def create_docx_bytes() -> bytes:
    buffer = BytesIO()
    document = Document()

    document.add_heading("Remote Work", level=1)
    document.add_paragraph("Employees may work remotely three days per week.")

    document.add_heading("Paid Leave", level=1)
    document.add_paragraph("Employees receive eighteen paid leave days.")

    document.save(buffer)

    return buffer.getvalue()


def test_ingest_document_processes_docx_with_heading_metadata() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "employee_handbook.docx",
                create_docx_bytes(),
                (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
        },
    )

    assert response.status_code == 200

    body = response.json()

    assert body["file_name"] == "employee_handbook.docx"
    assert body["file_extension"] == ".docx"
    assert body["chunk_count"] == 2

    assert body["chunks"][0]["section_index"] == 0
    assert body["chunks"][0]["page_number"] is None
    assert body["chunks"][0]["heading"] == "Remote Work"

    assert body["chunks"][1]["section_index"] == 1
    assert body["chunks"][1]["page_number"] is None
    assert body["chunks"][1]["heading"] == "Paid Leave"


def test_same_normalized_content_produces_same_document_id() -> None:
    first_response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "first.txt",
                b"Remote   work is allowed.\n",
                "text/plain",
            )
        },
    )

    second_response = client.post(
        "/api/v1/documents/ingest",
        files={
            "file": (
                "second.txt",
                b"Remote work is allowed.",
                "text/plain",
            )
        },
    )

    assert first_response.status_code == 200
    assert second_response.status_code == 200

    first_body = first_response.json()
    second_body = second_response.json()

    assert first_body["document_id"] == second_body["document_id"]
    assert first_body["content_hash"] == second_body["content_hash"]
    assert first_body["chunks"][0]["chunk_id"] == second_body["chunks"][0]["chunk_id"]
