from __future__ import annotations

import argparse
import json
import mimetypes
import sys
import tempfile
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen

TXT_MARKER = "ACCEPTANCE_TXT_MARKER_4f7a9c"
MD_MARKER = "ACCEPTANCE_MARKDOWN_MARKER_9d21c"
PDF_MARKER = "ACCEPTANCE_PDF_MARKER_62b4e"
DOCX_MARKER = "ACCEPTANCE_DOCX_DELETE_MARKER_0aa31"
REPLACEMENT_MARKER = "ACCEPTANCE_REPLACEMENT_MARKER_c81e2"
STALE_MARKER = "ACCEPTANCE_STALE_DOCX_MARKER_d72ab"
MISSING_MARKER = "ACCEPTANCE_NO_EVIDENCE_MARKER_6fb91"


class AcceptanceError(RuntimeError):
    pass


@dataclass
class ApiResponse:
    status: int
    data: dict[str, Any]
    request_id: str | None = None


@dataclass
class AcceptanceClient:
    base_url: str
    api_key: str | None = None
    timeout: float = 30.0
    auth_header_name: str = "X-API-Key"

    def headers(self, extra: dict[str, str] | None = None) -> dict[str, str]:
        headers = {"Accept": "application/json"}
        if self.api_key:
            headers[self.auth_header_name] = self.api_key
        if extra:
            headers.update(extra)
        return headers

    def json(
        self,
        method: str,
        path: str,
        payload: dict[str, Any] | None = None,
    ) -> ApiResponse:
        body = None
        headers = self.headers()
        if payload is not None:
            body = json.dumps(payload).encode("utf-8")
            headers["Content-Type"] = "application/json"
        return self._send(Request(self.base_url + path, body, headers, method=method))

    def upload(
        self,
        path: Path,
        *,
        replace_document_id: str | None = None,
    ) -> ApiResponse:
        boundary = f"----acceptance-{uuid.uuid4().hex}"
        fields: list[bytes] = []
        if replace_document_id:
            fields.extend(
                [
                    f"--{boundary}\r\n".encode(),
                    (
                        b"Content-Disposition: form-data; "
                        b'name="replace_document_id"\r\n\r\n'
                    ),
                    replace_document_id.encode(),
                    b"\r\n",
                ]
            )
        content_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        body = b"".join(
            [
                *fields,
                f"--{boundary}\r\n".encode(),
                (
                    'Content-Disposition: form-data; name="file"; '
                    f'filename="{path.name}"\r\n'
                ).encode(),
                f"Content-Type: {content_type}\r\n\r\n".encode(),
                path.read_bytes(),
                f"\r\n--{boundary}--\r\n".encode(),
            ]
        )
        headers = self.headers(
            {
                "Content-Type": f"multipart/form-data; boundary={boundary}",
                "Content-Length": str(len(body)),
            }
        )
        return self._send(
            Request(
                self.base_url + "/api/v1/documents/ingest",
                body,
                headers,
                method="POST",
            ),
            timeout=max(self.timeout, 180.0),
        )

    def _send(self, request: Request, *, timeout: float | None = None) -> ApiResponse:
        try:
            with urlopen(request, timeout=timeout or self.timeout) as response:
                return ApiResponse(
                    status=response.status,
                    data=parse_json(response.read()),
                    request_id=response.headers.get("X-Request-ID"),
                )
        except HTTPError as error:
            request_id = error.headers.get("X-Request-ID")
            detail = safe_error_detail(error.read())
            suffix = f" request_id={request_id}" if request_id else ""
            raise AcceptanceError(
                f"{request.get_method()} {request.full_url} HTTP {error.code}: "
                f"{detail}{suffix}"
            ) from error
        except URLError as error:
            raise AcceptanceError(
                f"{request.get_method()} {request.full_url} failed: {error.reason}"
            ) from error


@dataclass
class AcceptanceState:
    document_ids: set[str] = field(default_factory=set)
    checks: int = 0


def parse_json(body: bytes) -> dict[str, Any]:
    data = json.loads(body.decode("utf-8"))
    if not isinstance(data, dict):
        raise AcceptanceError("response was not a JSON object")
    return data


def safe_error_detail(body: bytes) -> str:
    try:
        data = json.loads(body.decode("utf-8", errors="replace"))
    except json.JSONDecodeError:
        return "non-json error response"
    detail = data.get("detail") if isinstance(data, dict) else None
    return str(detail)[:300] if detail is not None else "error response"


def pass_check(name: str, state: AcceptanceState) -> None:
    state.checks += 1
    print(f"[PASS] {name}")


def wait_for_liveness(client: AcceptanceClient, *, deadline_seconds: float) -> None:
    deadline = time.monotonic() + deadline_seconds
    last_error = "not attempted"
    while time.monotonic() < deadline:
        try:
            response = client.json("GET", "/api/v1/health/live")
            if response.status == 200 and response.data.get("status") == "alive":
                return
            last_error = f"unexpected liveness response: {response.data!r}"
        except AcceptanceError as error:
            last_error = str(error)
        time.sleep(1.0)
    raise AcceptanceError(f"API did not become live: {last_error}")


def assert_document_id(value: Any) -> str:
    if not isinstance(value, str) or len(value) != 64:
        raise AcceptanceError(f"invalid document_id: {value!r}")
    return value


def require_result_for(
    response: ApiResponse,
    *,
    document_id: str,
    marker: str,
) -> dict[str, Any]:
    results = response.data.get("results")
    if not isinstance(results, list):
        raise AcceptanceError("search results must be a list")
    for result in results:
        if result.get("document_id") == document_id and marker in result.get(
            "text", ""
        ):
            return result
    raise AcceptanceError(f"expected marker {marker} for document {document_id}")


def require_no_results(
    response: ApiResponse, *, document_id: str | None = None
) -> None:
    results = response.data.get("results")
    if not isinstance(results, list):
        raise AcceptanceError("search results must be a list")
    scoped = [
        result
        for result in results
        if document_id is None or result.get("document_id") == document_id
    ]
    if scoped:
        raise AcceptanceError(f"expected no scoped results, got {len(scoped)}")


def assert_safe_diagnostics(result: dict[str, Any]) -> None:
    diagnostics = result.get("score_diagnostics")
    if not isinstance(diagnostics, dict):
        raise AcceptanceError("score_diagnostics missing")
    for branch_name in ("dense", "sparse"):
        branch = diagnostics.get(branch_name)
        if not isinstance(branch, dict):
            raise AcceptanceError(f"{branch_name} diagnostics missing")
        for key in ("raw_score", "rank", "weight", "rrf_contribution"):
            if key not in branch:
                raise AcceptanceError(f"{branch_name}.{key} missing")
    if diagnostics.get("fused_rank", 0) < 1:
        raise AcceptanceError("fused_rank must be 1-based")
    if "fused_score" not in diagnostics:
        raise AcceptanceError("fused_score missing")
    forbidden = json.dumps(diagnostics).lower()
    for token in ("vector", "indices", "payload", "point_id"):
        if token in forbidden:
            raise AcceptanceError(f"diagnostics exposed {token}")


def project_root() -> Path:
    return Path(__file__).resolve().parents[2]


def fixture_dir() -> Path:
    return Path(__file__).resolve().parent / "fixtures"


def build_runtime_fixtures(tmp_dir: Path) -> dict[str, Path]:
    from docx import Document
    from reportlab.pdfgen import canvas

    pdf_path = tmp_dir / "acceptance_pdf.pdf"
    pdf = canvas.Canvas(str(pdf_path))
    pdf.drawString(72, 720, f"{PDF_MARKER} pdf reimbursement evidence.")
    pdf.save()

    docx_path = tmp_dir / "acceptance_delete.docx"
    doc = Document()
    doc.add_heading("Acceptance Delete Fixture", level=1)
    doc.add_paragraph(f"{DOCX_MARKER} delete target evidence. {STALE_MARKER}")
    doc.save(docx_path)

    return {
        "txt": fixture_dir() / "acceptance_policy.txt",
        "md": fixture_dir() / "acceptance_notes.md",
        "pdf": pdf_path,
        "docx": docx_path,
        "replacement": fixture_dir() / "replacement_policy.txt",
    }


def upload_all(
    client: AcceptanceClient,
    paths: dict[str, Path],
    state: AcceptanceState,
) -> dict[str, str]:
    ids = {}
    for label in ("txt", "md", "pdf", "docx"):
        response = client.upload(paths[label])
        document_id = assert_document_id(response.data.get("document_id"))
        ids[label] = document_id
        state.document_ids.add(document_id)
        pass_check(f"Uploaded {label.upper()}", state)
    return ids


def list_document_ids(client: AcceptanceClient) -> set[str]:
    response = client.json("GET", "/api/v1/documents?limit=100")
    documents = response.data.get("documents")
    if not isinstance(documents, list):
        raise AcceptanceError("documents must be a list")
    return {
        document["document_id"]
        for document in documents
        if isinstance(document.get("document_id"), str)
    }


def cleanup(client: AcceptanceClient, state: AcceptanceState) -> None:
    for document_id in sorted(state.document_ids):
        try:
            client.json("DELETE", f"/api/v1/documents/{document_id}")
        except AcceptanceError:
            pass


def run_workflow(client: AcceptanceClient, *, keep_data: bool) -> int:
    state = AcceptanceState()
    try:
        wait_for_liveness(client, deadline_seconds=60)
        pass_check("API liveness", state)
        try:
            client.json("GET", "/api/v1/health/ready")
        except AcceptanceError:
            pass

        with tempfile.TemporaryDirectory() as tmp:
            paths = build_runtime_fixtures(Path(tmp))
            ids = upload_all(client, paths, state)

            listed = list_document_ids(client)
            if not set(ids.values()) <= listed:
                raise AcceptanceError("uploaded documents missing from listing")
            pass_check("Document listing", state)

            detail = client.json("GET", f"/api/v1/documents/{ids['txt']}").data
            if detail.get("filename") != "acceptance_policy.txt":
                raise AcceptanceError("document detail filename mismatch")
            pass_check("Document detail", state)

            dense = client.json(
                "POST", "/api/v1/search/dense", {"query": TXT_MARKER, "limit": 5}
            )
            require_result_for(dense, document_id=ids["txt"], marker=TXT_MARKER)
            pass_check("Dense retrieval", state)

            sparse = client.json(
                "POST", "/api/v1/search/sparse", {"query": MD_MARKER, "limit": 5}
            )
            require_result_for(sparse, document_id=ids["md"], marker=MD_MARKER)
            pass_check("Sparse retrieval", state)

            hybrid = client.json(
                "POST", "/api/v1/search/hybrid", {"query": PDF_MARKER, "limit": 5}
            )
            require_result_for(hybrid, document_id=ids["pdf"], marker=PDF_MARKER)
            pass_check("Hybrid retrieval", state)

            diagnostics = client.json(
                "POST",
                "/api/v1/search/hybrid",
                {"query": TXT_MARKER, "limit": 3, "include_score_diagnostics": True},
            )
            diagnostic_result = require_result_for(
                diagnostics, document_id=ids["txt"], marker=TXT_MARKER
            )
            assert_safe_diagnostics(diagnostic_result)
            pass_check("Score diagnostics", state)

            filtered = client.json(
                "POST",
                "/api/v1/search/hybrid",
                {"query": TXT_MARKER, "document_id": ids["txt"], "limit": 5},
            )
            require_result_for(filtered, document_id=ids["txt"], marker=TXT_MARKER)
            if any(
                result.get("document_id") != ids["txt"]
                for result in filtered.data["results"]
            ):
                raise AcceptanceError("document_id filter returned another document")
            content_filtered = client.json(
                "POST",
                "/api/v1/search/hybrid",
                {"query": MD_MARKER, "content_types": ["text/markdown"], "limit": 5},
            )
            require_result_for(
                content_filtered, document_id=ids["md"], marker=MD_MARKER
            )
            pass_check("Metadata filters", state)

            replacement = client.upload(
                paths["replacement"], replace_document_id=ids["docx"]
            )
            new_doc_id = assert_document_id(replacement.data.get("document_id"))
            state.document_ids.add(new_doc_id)
            replaced = client.json(
                "POST",
                "/api/v1/search/hybrid",
                {"query": REPLACEMENT_MARKER, "document_id": new_doc_id, "limit": 5},
            )
            require_result_for(
                replaced, document_id=new_doc_id, marker=REPLACEMENT_MARKER
            )
            stale = client.json(
                "POST",
                "/api/v1/search/hybrid",
                {"query": STALE_MARKER, "document_id": ids["docx"], "limit": 5},
            )
            require_no_results(stale, document_id=ids["docx"])
            state.document_ids.discard(ids["docx"])
            ids["docx"] = new_doc_id
            pass_check("Document replacement", state)

            client.json("DELETE", f"/api/v1/documents/{ids['pdf']}")
            state.document_ids.discard(ids["pdf"])
            if ids["pdf"] in list_document_ids(client):
                raise AcceptanceError("deleted document remained in listing")
            deleted = client.json(
                "POST",
                "/api/v1/search/hybrid",
                {"query": PDF_MARKER, "document_id": ids["pdf"], "limit": 5},
            )
            require_no_results(deleted, document_id=ids["pdf"])
            pass_check("Document deletion", state)

            answer = client.json(
                "POST",
                "/api/v1/answers/grounded",
                {
                    "question": f"What evidence mentions {TXT_MARKER}?",
                    "document_id": ids["txt"],
                },
            ).data
            if not answer.get("citations") or answer.get("insufficient_context"):
                raise AcceptanceError("grounded answer did not cite evidence")
            if "score_diagnostics" in answer:
                raise AcceptanceError("grounded answer exposed diagnostics")
            insufficient = client.json(
                "POST",
                "/api/v1/answers/grounded",
                {"question": MISSING_MARKER, "document_id": "0" * 64},
            ).data
            if insufficient.get("finish_reason") != "insufficient_context":
                raise AcceptanceError("no-evidence answer was not insufficient")
            pass_check("Grounded answer", state)
    finally:
        if keep_data:
            print("[PASS] Cleanup skipped (--keep-data)")
        else:
            cleanup(client, state)
            pass_check("Cleanup", state)

    print(f"Acceptance workflow passed: {state.checks}/{state.checks} checks")
    return 0


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--base-url",
        default=None,
        help="API base URL. Defaults to ACCEPTANCE_API_BASE_URL or localhost.",
    )
    parser.add_argument("--api-key", default=None)
    parser.add_argument("--keep-data", action="store_true")
    parser.add_argument("--timeout", type=float, default=30.0)
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    import os

    args = parse_args(argv)
    client = AcceptanceClient(
        base_url=(
            args.base_url
            or os.getenv("ACCEPTANCE_API_BASE_URL")
            or "http://127.0.0.1:8000"
        ).rstrip("/"),
        api_key=args.api_key or os.getenv("ACCEPTANCE_API_KEY"),
        timeout=args.timeout,
    )
    try:
        return run_workflow(client, keep_data=args.keep_data)
    except Exception as error:
        print(f"[FAIL] {error}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
